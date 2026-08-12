import streamlit as st
import os
from io import BytesIO
from datetime import datetime

st.set_page_config(page_title="PDF FOTOS a WORD - PRO MAX", page_icon="📸", layout="wide")

st.markdown("""
<style>
    .main { background: #FAF9F5; }
    [data-testid="stSidebar"] { background: #F0EEE6; }
    .card { background: white; border-radius: 16px; padding: 20px; border: 1px solid #E8E5DD; box-shadow: 0 4px 12px rgba(0,0,0,0.06); margin-bottom:12px; }
    .title { font-weight:800; font-size:32px; background: linear-gradient(90deg,#D4A574,#8B5E34); -webkit-background-clip:text; -webkit-text-fill-color:transparent; }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="title">📸➡️📝 PDF con FOTOS a WORD - PRO MAX</div>', unsafe_allow_html=True)
st.caption("Fuerza máxima OCR - Convierte PDFs escaneados, fotos de libros, apuntes a mano")

# --- FUNCIONES OCR FORTALECIDAS ---
def ocr_image_pil(image_pil):
    """Intenta OCR con varios motores"""
    text = ""
    # Intento 1: pytesseract
    try:
        import pytesseract
        # Config para español + mejorar calidad
        custom_config = r'--oem 3 --psm 6 -l spa+eng'
        text = pytesseract.image_to_string(image_pil, config=custom_config)
        if len(text.strip()) > 20:
            return text
    except Exception as e:
        st.toast(f"pytesseract no disponible: {e}", icon="⚠️")
    
    # Intento 2: EasyOCR (funciona en Streamlit Cloud sin instalar nada extra)
    try:
        import easyocr
        # Cache del reader para no cargarlo cada vez
        if 'ocr_reader' not in st.session_state:
            with st.spinner("Cargando motor OCR (primera vez tarda 30s)..."):
                st.session_state.ocr_reader = easyocr.Reader(['es', 'en'], gpu=False)
        reader = st.session_state.ocr_reader
        results = reader.readtext(image_pil, detail=0, paragraph=True)
        text_easy = "\n".join(results)
        if len(text_easy.strip()) > len(text.strip()):
            text = text_easy
    except Exception as e:
        # st.toast(f"EasyOCR no disponible: {e}")
        pass
    
    return text

with st.sidebar:
    st.markdown("### ⚙️ Configuración OCR")
    dpi = st.slider("Calidad OCR (DPI)", 150, 400, 300, help="Más DPI = mejor lectura pero más lento")
    modo_fotos = st.radio("Qué hacer con fotos", ["📝 Extraer solo texto con OCR", "🖼️ Texto + guardar foto en WORD", "📄 Todo foto como imagen en WORD"], index=0)
    st.divider()
    st.markdown("**Requisitos para OCR perfecto:**")
    st.code("streamlit\nPyPDF2\npython-docx\nPillow\nPyMuPDF\neasyocr\npytesseract", language="text")

col1, col2 = st.columns([1.3, 0.9])

with col1:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    uploaded = st.file_uploader("📤 Sube PDF con fotos escaneadas", type=["pdf","png","jpg","jpeg"], accept_multiple_files=False)
    
    if uploaded:
        st.info(f"Archivo: {uploaded.name} - {uploaded.size/1024/1024:.2f} MB")
        
        if st.button("🚀 CONVERTIR FOTOS A WORD - FUERZA MÁXIMA", type="primary", use_container_width=True):
            from docx import Document
            from docx.shared import Inches
            from PIL import Image
            import PyPDF2
            
            doc = Document()
            doc.add_heading(f'Convertido: {uploaded.name}', 0)
            doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')} | Motor: OCR Doble (Tesseract + EasyOCR) | DPI: {dpi}")
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            full_text_all = ""
            
            # CASO 1: IMAGEN SUELTA
            if uploaded.type in ["image/png", "image/jpeg", "image/jpg"]:
                status_text.text("Procesando imagen suelta con OCR...")
                image = Image.open(uploaded).convert("RGB")
                text = ocr_image_pil(image)
                if not text.strip():
                    text = "[No se detectó texto en la imagen]"
                doc.add_heading("Texto extraído de la imagen", 1)
                doc.add_paragraph(text)
                if "guardar foto" in modo_fotos or "imagen" in modo_fotos:
                    # Guardar imagen en WORD
                    img_bio = BytesIO()
                    image.save(img_bio, format='PNG')
                    doc.add_picture(BytesIO(img_bio.getvalue()), width=Inches(5))
                full_text_all = text
                progress_bar.progress(1.0)
            
            # CASO 2: PDF
            else:
                try:
                    import fitz  # PyMuPDF
                    pdf_data = uploaded.getvalue()
                    doc_fitz = fitz.open(stream=pdf_data, filetype="pdf")
                    total_pages = len(doc_fitz)
                    
                    for i in range(total_pages):
                        status_text.text(f"Página {i+1}/{total_pages} - Analizando...")
                        progress_bar.progress((i+1)/total_pages)
                        
                        page = doc_fitz[i]
                        # 1. Intentar extraer texto digital primero
                        text_digital = page.get_text("text") or ""
                        
                        doc.add_heading(f"Página {i+1}", 2)
                        
                        if len(text_digital.strip()) > 100:
                            # Es PDF digital, usar texto directo
                            doc.add_paragraph(text_digital)
                            full_text_all += text_digital + "\n"
                        else:
                            # Es foto escaneada - HACER OCR
                            status_text.text(f"Página {i+1}/{total_pages} - Es foto, haciendo OCR (DPI {dpi})...")
                            pix = page.get_pixmap(dpi=dpi)
                            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                            
                            text_ocr = ocr_image_pil(img)
                            
                            if len(text_ocr.strip()) < 10:
                                text_ocr = f"[Página {i+1}: No se pudo leer texto. Puede ser foto muy borrosa o a mano muy desordenada]"
                            
                            doc.add_paragraph(text_ocr)
                            full_text_all += text_ocr + "\n"
                            
                            # Si el usuario quiere guardar fotos
                            if "guardar foto" in modo_fotos:
                                img_bio = BytesIO()
                                img.thumbnail((800, 800))
                                img.save(img_bio, format='PNG')
                                doc.add_picture(BytesIO(img_bio.getvalue()), width=Inches(5))
                        
                        # 2. Extraer imágenes embebidas del PDF también
                        try:
                            image_list = page.get_images(full=True)
                            if image_list and "guardar foto" in modo_fotos:
                                for img_index, img in enumerate(image_list[:2]): # max 2 por página
                                    xref = img[0]
                                    base_image = doc_fitz.extract_image(xref)
                                    image_bytes = base_image["image"]
                                    pil_img = Image.open(BytesIO(image_bytes))
                                    # OCR a esa imagen también
                                    text_from_embedded = ocr_image_pil(pil_img)
                                    if len(text_from_embedded.strip()) > 20:
                                        doc.add_paragraph(f"[Texto en imagen {img_index+1}]: {text_from_embedded}", style='Intense Quote')
                        except:
                            pass
                    
                    doc_fitz.close()
                except ImportError:
                    st.error("Falta PyMuPDF. Instala con: pip install PyMuPDF")
                except Exception as e:
                    st.error(f"Error procesando PDF: {e}")
            
            # Guardar WORD final
            bio = BytesIO()
            doc.save(bio)
            st.session_state['word_file'] = bio.getvalue()
            st.session_state['full_text'] = full_text_all
            st.session_state['orig_name'] = uploaded.name
            status_text.text("✅ ¡Conversión completada!")
            st.balloons()
    
    st.markdown('</div>', unsafe_allow_html=True)

with col2:
    st.markdown('<div class="card">', unsafe_allow_html=True)
    st.markdown("### 📥 Resultado")
    if 'word_file' in st.session_state:
        st.success(f"✅ Listo: {st.session_state.get('orig_name','')}")
        st.metric("Texto extraído", f"{len(st.session_state.get('full_text',''))} chars")
        
        st.download_button(
            "⬇️ DESCARGAR WORD FINAL",
            st.session_state['word_file'],
            file_name=f"{st.session_state.get('orig_name','doc').split('.')[0]}_OCR_WORD.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
        st.divider()
        st.markdown("**Vista previa texto OCR:**")
        st.text_area("Texto", st.session_state.get('full_text','')[:4000], height=400)
    else:
        st.markdown("""
        **Esta versión PRO MAX hace:**
        
        1. **Detecta automáticamente** si cada página es texto digital o foto
        2. **Si es foto:** la convierte a imagen HD y le hace OCR doble
        3. **Doble motor:** Primero Tesseract, si falla usa EasyOCR (funciona en la nube)
        4. **Extrae texto de imágenes dentro del PDF** también
        5. **Soporta 1000 páginas** con barra de progreso
        
        **¿No lee fotos aún?** 
        En Streamlit Cloud necesitas agregar en `packages.txt`:
        ```
        tesseract-ocr
        tesseract-ocr-spa
        ```
        """)
    st.markdown('</div>', unsafe_allow_html=True)

st.divider()
st.markdown("### 📋 Actualiza tu requirements.txt y packages.txt")
colA, colB = st.columns(2)
with colA:
    st.markdown("**requirements.txt**")
    st.code("streamlit\nPyPDF2\npython-docx\nPillow\nPyMuPDF\neasyocr\npytesseract", language="text")
with colB:
    st.markdown("**packages.txt (archivo nuevo en GitHub)**")
    st.code("tesseract-ocr\ntesseract-ocr-spa\nlibgl1", language="text")
    st.caption("Crea este archivo nuevo en GitHub junto a app.py para que el OCR funcione en la nube")

 
